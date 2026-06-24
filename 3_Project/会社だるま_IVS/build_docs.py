# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CHERRY = RGBColor(0xB1, 0x12, 0x26)
NAVY = RGBColor(0x24, 0x30, 0x5E)
INK = RGBColor(0x2A, 0x14, 0x16)

def setup_styles(doc):
    n = doc.styles['Normal']
    n.font.name = 'Hiragino Sans'
    n.font.size = Pt(10.5)

def add_title(doc, text, color):
    h = doc.add_heading(text, level=0)
    for r in h.runs:
        r.font.color.rgb = color
        r.font.name = 'Hiragino Sans'

def add_h(doc, text, level, color):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = color
        r.font.name = 'Hiragino Sans'

def meta(doc, lines):
    for label, val in lines:
        p = doc.add_paragraph()
        rb = p.add_run(label)
        rb.bold = True
        p.add_run(val)

def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)

def numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Number')
        if isinstance(it, tuple):
            r = p.add_run(it[0]); r.bold = True
            p.add_run(it[1])
        else:
            p.add_run(it)

def checklist(doc, items):
    for it in items:
        doc.add_paragraph('☐ ' + it, style='List Bullet')

def table(doc, headers, rows, accent):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h); r.bold = True
        r.font.color.rgb = accent
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9.5)
    return t

# ============ CORPORATE ============
def build_corporate(path):
    doc = Document()
    setup_styles(doc)
    add_title(doc, '会社だるま × IVS｜企業納品動画 企画書', CHERRY)
    meta(doc, [
        ('目的：', 'IVS会場で撮影する「企業納品動画」の狙いと構成を整理し、代表確認のうえ協力企業に共有できる状態にする。'),
        ('作成：', 'SNSプランナー　／　ステータス：代表確認待ち（v1.0）'),
    ])
    p = doc.add_paragraph()
    r = p.add_run('※ 本企画は IVSイベント会場での納品動画。デスク設置カットなど“オフィス納品編”は第2弾として別途設計する。')
    r.italic = True

    add_h(doc, '1. 目的', 1, CHERRY)
    doc.add_paragraph('購入企業（既存顧客）向け。顧客満足＋法人による二次拡散（BtoB）→ 次の受注が狙い。')
    numbered(doc, [
        ('納品物としての価値：', '購入企業に「特別な体験の記録」を残し、サービス満足度を高める'),
        ('企業による二次拡散：', '企業自身がSNS投稿＝取引先・同業など“法人ネットワーク”へ届く'),
        ('新規法人への訴求：', '「自社もやりたい」を生み、次の受注につなげる'),
    ])

    add_h(doc, '2. 基本設計', 1, CHERRY)
    bullets(doc, [
        ('配信先：', 'Instagram Reels / TikTok / YouTube Shorts → 縦型 9:16'),
        ('想定尺：', '30〜60秒（フックは冒頭3秒で完結させる）'),
        ('機材：', 'スマホ＋ジンバル／手元用サブ1台／ピンマイク推奨'),
        ('演出の核：', '目入れ＝願掛け・決意宣言。願いを描く瞬間をクライマックスに'),
        ('感情設計：', '開封のワクワク → 目入れの集中 → 宣言の表情、で起伏をつくる'),
    ])
    add_h(doc, '撮影・編集の型（+α 改善ポイント）', 2, INK)
    bullets(doc, [
        ('音設計（ASMR）：', '開封音・ペンの目入れ音を意図的に録音。「音ON推奨」テロップで音声視聴を誘発'),
        ('構図・ループ：', '重要な文字／顔は画面中央60%（縦型セーフゾーン）に。ラスト→冒頭でシームレスループ＝再生回数UP'),
        ('テロップ2種：', '質問/状況テロップ＋ネームスーパー（社名・氏名）で無音視聴に対応＝完視聴率UP'),
    ])

    add_h(doc, '3. 構成（カット表）', 1, CHERRY)
    table(doc, ['#', 'カット', '尺目安', '内容・狙い', 'ディレクション'], [
        ['1', 'フック', '0–3秒', '「頑張りたいこと」＋会社名を本人が宣言', '最初の一言で掴む。例「○○を実現します、株式会社△△です」'],
        ['2', '開封', '3–8秒', '箱を開ける様子', '手元アップ＋開封音をしっかり録る（音ON推奨）'],
        ['3', '取り出し（ヒキ）', '8–13秒', '箱から取り出して驚く様子', '引きで全身＋表情。リアクションを大切に'],
        ['4', '目標を書く（クローズ）', '13–19秒', 'だるまに願い・社名を書く手元アップ', '会社だるま最大の差別化＝必ず魅せる。ペンの音も録る'],
        ['5', '目入れPOV', '19–26秒', 'ペンが目に迫る主観カット', '筆／ペン先のアップ。緊張感のあるクライマックス'],
        ['6', '目入れ（ヒキ）', '26–32秒', '目を入れる瞬間を引きで', '表情と手元が両方入る画角'],
        ['7', '意気込み', '32–45秒', '決意表明で締め', 'カメラ目線で力強く。社名ロゴ／だるまで締め'],
    ], CHERRY)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('ディレクションの軸').bold = True
    bullets(doc, [
        '冒頭の宣言を締めの意気込みで言い換え、一本の物語として繋げる',
        '開封の“素のリアクション”が命。事前に中身を見せすぎない',
        '目標を書く手元と目入れ（POV＋ヒキ）は編集の山場＝必ず押さえる',
    ])

    add_h(doc, '4. 撮影時のお願い（協力企業向け）', 1, CHERRY)
    add_h(doc, '事前のご準備', 2, INK)
    bullets(doc, [
        '冒頭で言う「目標（頑張りたいこと）」＋「会社名」を1文でご準備ください',
        'だるまに書く願い・社名の文言を事前に決めておいてください',
        '撮影に出る方（顔出し）を1〜数名お決めください',
    ])
    add_h(doc, '当日のお願い', 2, INK)
    bullets(doc, [
        '開封〜目入れは“素のリアクション”が魅力。リハなし・一発本番でOK',
        '撮影時間：1社あたり 約◯分（要確定）',
        '顔出し・SNS掲載可否の同意をお願いします',
    ])

    add_h(doc, '5. 当日オペレーション＆準備物', 1, CHERRY)
    p = doc.add_paragraph(); p.add_run('オペレーション').bold = True
    bullets(doc, [
        '役割分担：撮影／進行／同意・誘導 を分担',
        '撮影ブース：ロゴが映える背景、開封スペース、目入れ用の机',
        'だるまは事前に準備（社名・目標を「書く工程」の有無を確認）',
    ])
    p = doc.add_paragraph(); p.add_run('準備物チェックリスト').bold = True
    checklist(doc, [
        'だるま（社名・目標 記入用）', '目入れ用ペン・筆・墨', '開封演出用の箱',
        'ジンバル・スマホ（メイン／サブ）', 'ピンマイク・モバイルバッテリー', '背景ボード／ロゴ',
        '同意確認フォーム（掲載可否）',
    ])

    add_h(doc, '6. 編集・納品', 1, CHERRY)
    bullets(doc, [
        ('尺・縦横：', '30〜60秒／縦型 9:16'),
        ('テロップ：', '社名・目標・キーワードを最小限。だるま文化の一言を統一フォーマットで'),
        ('BGM：', '高揚感のある共通BGMでシリーズ感（権利フリー音源）'),
        ('納品形式：', '各社へ個別データを納品（縦型／必要に応じ横型も）'),
        ('資産化：', 'OP/EDのロゴ演出を共通化し、ブランド資産として蓄積'),
    ])

    add_h(doc, '7. 代表に確認したいこと（決め事）', 1, CHERRY)
    numbered(doc, [
        ('配信先：', 'Reels／TikTok／Shorts の優先順位は？'),
        ('撮影社数：', '企業納品は何社を目標にするか'),
        ('撮影時間：', '1社あたりの上限の目安（オペ設計に必要）'),
        ('同意・二次利用：', '掲載可否の取り方と、納品データの二次利用ルール'),
        ('ブランド素材：', 'ロゴ・指定カラー・指定フォント・使用可能BGMの有無'),
        ('トーンの方向性：', '「かっこいい/エモい」寄りか「親しみ/応援」寄りか'),
    ])
    doc.save(path)
    print('Saved', path)

# ============ INTERVIEW ============
def build_interview(path):
    doc = Document()
    setup_styles(doc)
    add_title(doc, '会社だるま × IVS｜参加者インタビュー動画 企画書', NAVY)
    meta(doc, [
        ('目的：', 'IVS会場で撮影する「参加者インタビュー動画」の狙いと構成を整理し、代表確認のうえ運用できる状態にする。'),
        ('作成：', 'SNSプランナー　／　ステータス：代表確認待ち（v1.0）'),
    ])

    add_h(doc, '1. 目的', 1, NAVY)
    doc.add_paragraph('一般参加者向け。一般層への認知・共感（BtoC）→ フォロワー獲得が狙い。')
    numbered(doc, [
        ('会場の熱量を発信：', '参加者の生の声で、イベントとブランドの世界観を届ける'),
        ('共感でUGC化：', '「叶えたいこと」という普遍テーマで一般視聴者の共感・保存・シェアを誘発'),
        ('新規フォロワー獲得：', '公式アカウントの認知を一般層（＝将来の見込み客）へ拡大'),
    ])

    add_h(doc, '2. 基本設計', 1, NAVY)
    bullets(doc, [
        ('配信先：', 'Instagram Reels / TikTok / YouTube Shorts → 縦型 9:16'),
        ('想定尺：', '30〜60秒（フックは冒頭3秒で完結させる）'),
        ('機材：', 'スマホ＋ジンバル／手元用サブ1台／ピンマイク推奨'),
        ('演出の核：', '目入れ＝願掛け・決意宣言。願いを描く瞬間をクライマックスに'),
        ('感情設計：', '自己紹介で親しみ → 「叶えたいこと」で共感 → 目入れの表情で締める'),
    ])
    add_h(doc, '撮影・編集の型（+α 改善ポイント）', 2, INK)
    bullets(doc, [
        ('音設計（ASMR）：', 'ペンの目入れ音などを意図的に録音。「音ON推奨」テロップで音声視聴を誘発'),
        ('構図・ループ：', '重要な文字／顔は画面中央60%（縦型セーフゾーン）に。ラスト→冒頭でシームレスループ＝再生回数UP'),
        ('テロップ2種：', '質問テロップ＋ネームスーパー（名前・肩書き）で無音視聴に対応＝完視聴率UP'),
    ])

    add_h(doc, '3. 構成（質問フロー）', 1, NAVY)
    table(doc, ['#', 'パート', '質問・内容', '狙い'], [
        ['0', 'フック（先出し）', '一番エモい「叶えたいこと」を冒頭に先出し → 自己紹介へ繋ぐ', '完視聴率を上げる構成。スクロールを止める'],
        ['1', '自己紹介', 'お名前・ステータス（学生／職業など）', '視聴者が人物像を掴む'],
        ['2', '参加理由', '「IVSに参加した理由は？」', '会場のリアル・熱量'],
        ['3', 'マイブーム', '「最近ハマっていることは？」（例：ランニング、Setlog 等）', '親しみ・人柄を出す'],
        ['4', '叶えたいこと', '「あなたが叶えたいことを教えてください」（留学／起業／資格取得／一人旅 など）', '動画の核。感情の山'],
        ['5', '目入れ＆記念撮影', '願いを込めて目入れ → 記念写真', 'だるま文化の体験＋締めの画'],
    ], NAVY)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.add_run('ディレクションの軸').bold = True
    bullets(doc, [
        '「叶えたいこと」が主役。話しやすいよう聞き手が相槌でリード',
        '発言の直後に目入れする流れで感情を乗せる',
        '記念写真は縦・横どちらも（本人へのプレゼント＋投稿素材）',
        'テンポよく、1人あたり短時間で回せる質問順にしている',
    ])

    add_h(doc, '4. 撮影時のお願い（参加者向け）', 1, NAVY)
    add_h(doc, '事前のご準備', 2, INK)
    bullets(doc, [
        '「叶えたいこと」を一言で考えておいてください（動画の主役になります）',
        '自己紹介（名前・ステータス）とマイブームを軽くご準備',
    ])
    add_h(doc, '当日のお願い', 2, INK)
    bullets(doc, [
        'テンポよく短時間で撮影します（1人あたり 約◯分・要確定）',
        '顔出し・SNS掲載可否の同意をお願いします',
        '撮影後、記念写真データをプレゼントします',
    ])

    add_h(doc, '5. 当日オペレーション＆準備物', 1, NAVY)
    p = doc.add_paragraph(); p.add_run('オペレーション').bold = True
    bullets(doc, [
        '役割分担：撮影／聞き手（インタビュアー）／同意・誘導 を分担',
        '撮影ブース：ブランドが映える背景・ロゴ、目入れ用の机、記念写真スポット',
        '記念写真の受け渡し：その場AirDrop／QR／後日送付 のいずれか',
    ])
    p = doc.add_paragraph(); p.add_run('準備物チェックリスト').bold = True
    checklist(doc, [
        'だるま（体験用）', '目入れ用ペン・筆・墨', 'ジンバル・スマホ（メイン／サブ）',
        'ピンマイク・モバイルバッテリー', '背景ボード／ロゴ', '同意確認フォーム（掲載可否）',
        '記念写真の受け渡し手段（QR等）',
    ])

    add_h(doc, '6. 編集・納品', 1, NAVY)
    bullets(doc, [
        ('尺・縦横：', '30〜60秒／縦型 9:16'),
        ('テロップ：', '質問・名前・肩書きを最小限。だるま文化の一言を統一フォーマットで'),
        ('BGM：', '高揚感のある共通BGMでシリーズ感（権利フリー音源）'),
        ('納品形式：', '会社だるま公式アカウントで投稿。複数人の“まとめ版”も展開可'),
        ('資産化：', 'OP/EDのロゴ演出を共通化し、ブランド資産として蓄積'),
    ])

    add_h(doc, '7. 代表に確認したいこと（決め事）', 1, NAVY)
    numbered(doc, [
        ('配信先：', 'Reels／TikTok／Shorts の優先順位は？'),
        ('撮影人数：', 'インタビューは何人を目標にするか'),
        ('撮影時間：', '1人あたりの上限の目安（オペ設計に必要）'),
        ('同意・肖像権：', '掲載可否の取り方（口頭／書面）'),
        ('ブランド素材：', 'ロゴ・指定カラー・指定フォント・使用可能BGMの有無'),
        ('トーンの方向性：', '「かっこいい/エモい」寄りか「親しみ/応援」寄りか'),
    ])
    doc.save(path)
    print('Saved', path)

BASE = '/Users/aoi/Desktop/Bluenote/3_Project/会社だるま_IVS/'
build_corporate(BASE + '企業納品動画_企画書.docx')
build_interview(BASE + 'インタビュー動画_企画書.docx')
